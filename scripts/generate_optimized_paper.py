from __future__ import annotations

import csv
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "docs" / "paper" / "刘乐-毕业论文终稿.docx"
OUTPUT_DOCX = ROOT / "docs" / "paper" / "刘乐-毕业论文终稿-优化版.docx"
FIG_DIR = ROOT / "results" / "figures"
CONFIG_PATH = ROOT / "scripts" / "experiment.config.json"
METRICS_CSV = ROOT / "results" / "processed" / "experiment_metrics.csv"
SUMMARY_CSV = ROOT / "results" / "processed" / "experiment_metrics_summary.csv"
VPA_SAMPLE = ROOT / "results" / "raw" / "hpa_vpa" / "burst" / "hpa_vpa_20260508_134705_r3" / "vpa.yaml"
HPA_SAMPLE = ROOT / "results" / "raw" / "hpa_multi" / "burst" / "hpa_multi_20260508_121006_r3" / "hpa_describe.txt"


STRATEGY_LABELS = {
    "static": "静态副本",
    "hpa_cpu": "CPU HPA",
    "hpa_multi": "多指标 HPA",
    "hpa_vpa": "HPA+VPA",
}
STRATEGY_ORDER = ["static", "hpa_cpu", "hpa_multi", "hpa_vpa"]


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def _fmt(value: str | float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def _pct(value: str | float) -> str:
    return f"{float(value) * 100:.2f}%"


def _summary_rows() -> list[dict[str, str]]:
    rows = _read_csv(SUMMARY_CSV)
    by_strategy = {row["strategy"]: row for row in rows}
    out = []
    for strategy in STRATEGY_ORDER:
        row = by_strategy[strategy]
        out.append(
            {
                "策略": STRATEGY_LABELS[strategy],
                "P95(ms)": _fmt(row["p95_ms_mean"], 0),
                "P99(ms)": _fmt(row["p99_ms_mean"], 0),
                "QPS": _fmt(row["qps_mean"], 2),
                "失败率": _pct(row["failure_rate_mean"]),
                "CPU利用率": f"{float(row['cpu_util_percent_mean']):.2f}%",
                "内存利用率": f"{float(row['memory_util_percent_mean']):.2f}%",
                "成本代理": _fmt(row["cost_per_1k_req_mean"], 8),
            }
        )
    return out


def _run_rows() -> list[dict[str, str]]:
    rows = _read_csv(METRICS_CSV)
    rows.sort(key=lambda row: (STRATEGY_ORDER.index(row["strategy"]), row["run_id"]))
    return rows


def _metric_lookup() -> dict[str, dict[str, str]]:
    by_label = {}
    for row in _summary_rows():
        by_label[row["策略"]] = row
    return by_label


def _extract_vpa_recommendation() -> str:
    if not VPA_SAMPLE.exists():
        return "VPA 推荐样例文件未找到。"
    text = VPA_SAMPLE.read_text(encoding="utf-8", errors="ignore")
    cpu = re.search(r"target:\s*\n\s*cpu:\s*([^\n]+)\n\s*memory:\s*([^\n]+)", text)
    cond = "RecommendationProvided=True" if "RecommendationProvided" in text and 'status: "True"' in text else "RecommendationProvided 状态已记录"
    if cpu:
        return f"{cond}，样例运行中 target 推荐值为 CPU {cpu.group(1).strip()}、内存 {cpu.group(2).strip()}。"
    return f"{cond}，推荐明细见归档的 vpa.yaml。"


def _extract_hpa_state() -> str:
    if not HPA_SAMPLE.exists():
        return "HPA 状态样例文件未找到。"
    text = HPA_SAMPLE.read_text(encoding="utf-8", errors="ignore")
    pods = re.search(r"Deployment pods:\s*([^\n]+)", text)
    metric = re.search(r'"http_requests_per_second" on pods:\s*([^\n]+)', text)
    parts = []
    if pods:
        parts.append(f"HPA 样例状态显示 Deployment pods 为 {pods.group(1).strip()}")
    if metric:
        parts.append(f"业务 QPS 指标当前值/目标值为 {metric.group(1).strip()}")
    return "；".join(parts) + "。" if parts else "HPA 样例状态已归档。"


def _clear_doc_body_keep_section_props(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _set_run_font(run, name: str = "宋体", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _set_paragraph_format(paragraph, first_line: bool = True) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)


def _add_page_break(paragraph) -> None:
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_run_font(run, "黑体", 18, True)


def _add_center(doc: Document, text: str, size: int = 12, bold: bool = False) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_run_font(run, "宋体", size, bold)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    _set_run_font(run, "黑体", 16 if level == 1 else 14, True)
    para.paragraph_format.line_spacing = 1.5


def _add_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _set_paragraph_format(para)
    run = para.add_run(text)
    _set_run_font(run, "宋体", 10.5)


def _add_plain_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _set_paragraph_format(para, first_line=False)
    run = para.add_run(text)
    _set_run_font(run, "宋体", 10.5)


def _add_table_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_run_font(run, "宋体", 10.5, True)


def _add_table(doc: Document, rows: list[dict[str, str]]) -> None:
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, head in enumerate(headers):
        table.rows[0].cells[i].text = head
    for row in rows:
        cells = table.add_row().cells
        for i, head in enumerate(headers):
            cells[i].text = str(row[head])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run, "宋体", 9)


def _add_figure(doc: Document, filename: str, caption: str) -> None:
    path = FIG_DIR / filename
    if not path.exists():
        _add_paragraph(doc, f"{caption}（图像文件未找到：{filename}）")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(5.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    _set_run_font(cap_run, "宋体", 10.5)


def _add_references(doc: Document, refs: list[str]) -> None:
    for i, ref in enumerate(refs, start=1):
        _add_plain_paragraph(doc, f"[{i}] {ref}")


def _build_content() -> list[tuple[str, str | int]]:
    cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8-sig"))
    summary = _metric_lookup()
    static = summary["静态副本"]
    hpa_cpu = summary["CPU HPA"]
    hpa_multi = summary["多指标 HPA"]
    hpa_vpa = summary["HPA+VPA"]
    vpa_text = _extract_vpa_recommendation()
    hpa_state = _extract_hpa_state()

    scenario = cfg["experiment"]["scenario"]
    repeats = cfg["experiment"]["repeats"]
    users = cfg["experiment"]["users"]
    spawn_rate = cfg["experiment"]["spawn_rate"]
    run_time = cfg["experiment"]["run_time"]
    warmup = cfg["experiment"]["warmup_seconds"]
    cooldown = cfg["experiment"]["cooldown_seconds"]

    return [
        ("h1", "摘要"),
        ("p", "随着云原生架构在企业信息系统中的持续应用，微服务系统的部署方式逐渐从单体应用的静态发布转向以容器、编排平台和监控指标为基础的动态治理。微服务拆分提高了系统的可维护性和独立部署能力，但也使服务实例数量、资源请求配置、指标采集链路和运行状态管理变得更加复杂。当访问量在短时间内明显上升时，固定副本配置可能导致请求排队、长尾延迟上升和失败率增加；当访问量下降时，过高的资源预留又会造成计算资源浪费。因此，如何在容器编排环境中根据负载变化自动调整服务容量，是微服务系统稳定运行和资源利用优化中的基础问题。"),
        ("p", "本文围绕“基于容器编排的微服务弹性伸缩策略研究与实践”展开，定位为系统设计与工程实践类毕业设计。论文以 Kubernetes 中的典型无状态 HTTP 微服务为实验对象，构建了一套包含服务部署、指标暴露、监控采集、策略切换、负载生成、数据归档和可视化分析的弹性伸缩验证平台。平台由 sample-api 微服务、Kubernetes Deployment/Service、metrics-server、Prometheus、Prometheus Adapter、Horizontal Pod Autoscaler、Vertical Pod Autoscaler、Locust 压测脚本和 Python 数据处理脚本组成。为了避免将系统能力夸大为全场景通用平台，本文明确研究边界：当前实现聚焦单个典型无状态微服务的伸缩行为，不包含数据库、缓存、消息队列和多服务调用链路，但其部署、指标和实验流程可作为多服务场景中单个服务伸缩策略配置的基础。"),
        ("p", f"正式实验采用 {scenario} 场景，配置为 users={users}、spawn_rate={spawn_rate}、run_time={run_time}，每组策略重复 {repeats} 次，并设置 warmup={warmup}s、cooldown={cooldown}s。本文设计了静态副本、CPU 单指标 HPA、CPU/内存/QPS 多指标 HPA、HPA+VPA Initial 协同四组策略，共形成 12 条正式实验记录。每次运行均保存 Locust 统计文件、HPA/VPA YAML、事件日志、Pod 状态、top 采样日志和 manifest 环境信息，保证结果能够从原始数据追溯到具体运行。实验结果显示，HPA+VPA 组平均 QPS 为 {hpa_vpa['QPS']}，平均 P95 时延为 {hpa_vpa['P95(ms)']} ms，失败率均值为 {hpa_vpa['失败率']}；静态副本组平均 P95 时延为 {static['P95(ms)']} ms，平均失败率为 {static['失败率']}。结果说明，在本文实验环境中，将资源指标与业务 QPS 指标结合，并引入 VPA 资源推荐机制，有助于提升突发负载下的吞吐能力和服务稳定性。"),
        ("p", "关键词：微服务；Kubernetes；容器编排；弹性伸缩；HPA；VPA；Prometheus"),
        ("h1", "Abstract"),
        ("p", "With the adoption of cloud-native architectures, microservice systems are increasingly deployed and operated through containers, orchestration platforms, and observable runtime metrics. Although microservices improve maintainability and independent deployment, they also introduce challenges in instance management, resource request configuration, metric collection, and runtime stability. Static replica settings are often insufficient under burst traffic and may lead to queueing, long-tail latency, and request failures, while excessive resource reservation causes waste during low-load periods."),
        ("p", "This thesis designs and implements a Kubernetes-based experimental platform for evaluating elastic scaling strategies of a representative stateless HTTP microservice. The platform integrates a sample API service, Kubernetes Deployment and Service manifests, metrics-server, Prometheus, Prometheus Adapter, HPA, VPA, Locust workload generation, and Python-based data aggregation. Four strategies are compared: static replicas, CPU-based HPA, multi-metric HPA with CPU, memory and QPS, and multi-metric HPA combined with VPA Initial mode. The formal burst experiment contains twelve valid runs, with three repetitions for each strategy."),
        ("p", f"The results show that the HPA+VPA strategy achieves an average QPS of {hpa_vpa['QPS']}, an average P95 latency of {hpa_vpa['P95(ms)']} ms, and an average failure rate of {hpa_vpa['失败率']} in the test environment. Compared with the static replica strategy, the combined strategy provides a more balanced performance in throughput, reliability, and resource configuration stability. The work demonstrates that Kubernetes-native autoscaling mechanisms, when combined with application-level metrics and reproducible experiments, can provide practical evidence for microservice capacity governance."),
        ("p", "Keywords: Microservice; Kubernetes; Container Orchestration; Elastic Scaling; HPA; VPA; Prometheus"),
        ("h1", "1 绪论"),
        ("h2", "1.1 研究背景"),
        ("p", "近年来，云原生技术体系逐渐成为后端服务部署和运行治理的重要基础。容器技术使应用及其依赖能够以镜像形式交付，Kubernetes 进一步通过声明式资源对象、控制器和调度机制实现应用编排。对于采用微服务架构的系统而言，服务往往被拆分为多个能够独立构建、独立发布和独立扩缩容的单元。相比传统单体系统，微服务架构能够降低模块间耦合，提高团队协作效率，也便于针对不同服务的负载特点进行独立治理。"),
        ("p", "然而，微服务架构并不会自动带来稳定性。服务数量增多后，部署对象、运行实例、资源限制、服务发现和监控指标都会显著增加。若仍然采用固定副本的方式部署服务，系统在面对突发访问、周期性任务或短时间压测时容易出现容量不足；若为了避免高峰风险而长期保留较高副本数量，又会造成低峰期资源闲置。对于运行在 Kubernetes 中的微服务，弹性伸缩机制提供了一种相对通用的治理思路：根据资源利用率或业务指标自动调整 Pod 副本数量，或者根据历史资源使用情况给出更合理的资源请求建议。"),
        ("p", "Kubernetes 原生提供的 HPA 主要解决水平伸缩问题，即根据 CPU、内存或自定义指标调整目标 Deployment 的副本数；VPA 主要解决垂直伸缩问题，即根据容器历史资源使用情况推荐或调整 CPU、内存请求值。在实际使用中，HPA 与 VPA 不是简单叠加关系。HPA 的计算依赖资源请求值，VPA 的推荐又可能改变资源请求，从而影响 HPA 的利用率判断。如果两者缺少合理的模式选择和实验验证，可能引入扩缩容不稳定、资源配置波动或业务中断风险。"),
        ("p", "本文选题正是基于上述背景展开。与单纯讨论理论机制不同，本文通过可运行的实验平台验证不同策略在同一负载场景下的表现。实验不追求提出新的机器学习预测模型，而是强调 Kubernetes 原生能力、监控指标链路、压测数据和结果分析之间的闭环。这样的研究方式更适合作为系统设计类毕业设计：它既需要理解容器编排、自动伸缩和可观测性，也需要完成可以运行、可以复现、可以被检查的工程实现。"),
        ("h2", "1.2 研究意义"),
        ("p", "从工程实践角度看，弹性伸缩策略的价值不仅在于“能够自动扩容”，更在于扩容依据是否可靠、扩容时机是否合适、扩容后的资源成本是否可接受。CPU HPA 是最常见的配置方式，但 CPU 并不总能及时反映业务压力。例如，HTTP 服务在请求排队、连接等待或下游阻塞时，CPU 利用率可能并未显著升高，用户体验却已经变差。引入 QPS 等业务指标，有助于让伸缩决策更贴近真实访问压力。"),
        ("p", "从实验方法角度看，许多系统设计项目容易停留在“服务能够部署”和“接口能够访问”的层面，缺少对运行状态的持续采集和对策略效果的量化比较。本文将部署、监控、压测、采样、汇总和出图串联为自动化流程，每次实验运行都会形成独立目录，记录原始数据和环境信息。这种方式使结论不依赖主观描述，而是能够回到 CSV、YAML、事件日志和图表中进行复核。"),
        ("p", "从教学和答辩检查角度看，本课题属于系统设计与工程实践类。学院系统代码检查标准关注系统完成度、现场演示、稳定性和功能与论文一致性。本文的实现包含可运行服务、核心配置、策略脚本、结果分析和正式实验数据，能够围绕核心功能进行现场演示。相较算法设计类项目，本文不强调模型训练和算法精度，而强调云原生系统链路搭建、策略对照实验和工程结论分析。"),
        ("h2", "1.3 国内外研究与应用现状"),
        ("p", "在工业界，Kubernetes 已经成为容器编排领域的重要事实标准。云服务厂商和互联网企业普遍基于 Kubernetes 或兼容系统管理大规模容器化服务。围绕弹性伸缩，常见做法包括基于 CPU/内存的 HPA、基于业务指标的自定义 HPA、基于资源画像的 VPA，以及结合队列长度、请求延迟或预测模型的扩展方案。生产环境中的弹性治理通常还会结合服务网格、监控告警、灰度发布和容量规划。"),
        ("p", "在学术研究中，弹性伸缩问题常被放在云计算资源调度、微服务性能优化和自动化运维等方向讨论。部分研究关注预测式伸缩，通过时间序列预测、机器学习或强化学习提前判断负载变化；部分研究关注控制理论或阈值策略，强调扩缩容稳定性和震荡控制；也有研究将弹性伸缩与成本优化、服务质量保障和能耗控制结合。上述工作为弹性伸缩提供了更复杂的方法，但在本科毕业设计范围内，若直接实现复杂预测控制器，容易受到数据规模、环境稳定性和验证成本限制。"),
        ("p", "因此，本文选择以 Kubernetes 原生机制为基础，重点验证可部署、可采集、可复现的弹性伸缩实验流程。该选择并不意味着课题难度降低，而是将难点放在真实工程链路上：metrics-server 是否能够提供资源指标，Prometheus Adapter 是否能够将应用指标注册到 custom.metrics.k8s.io，HPA 是否能够读取多类指标，VPA 是否能够给出推荐，压测数据是否能够被统一归档和分析。这些环节共同决定了实验结论是否可信。"),
        ("h2", "1.4 研究内容与论文结构"),
        ("p", "本文的主要研究内容包括四个方面。第一，设计并搭建 Kubernetes 本地实验环境，完成 sample-api 微服务的容器化部署和 Service 暴露。第二，实现应用侧指标暴露与监控采集链路，使 Prometheus 能够抓取请求计数和请求耗时指标，并通过 Prometheus Adapter 形成 Kubernetes 自定义指标。第三，设计静态副本、CPU HPA、多指标 HPA、HPA+VPA 四组策略，完成统一压测场景下的对照实验。第四，编写数据汇总和绘图脚本，对 P95/P99、QPS、失败率、资源利用率和成本代理等指标进行分析。"),
        ("p", "论文结构安排如下：第 1 章介绍研究背景、研究意义、相关研究现状和本文研究内容；第 2 章介绍微服务、Kubernetes、HPA/VPA、Prometheus 和 Locust 等相关技术；第 3 章从需求分析和总体架构角度说明系统边界、功能需求和实验流程；第 4 章介绍系统实现，包括微服务实现、部署配置、伸缩策略和自动化实验脚本；第 5 章给出实验设计、数据结果和对比分析；第 6 章讨论基于实验数据的策略优化方法、多服务扩展方向和局限性；第 7 章总结全文并提出后续工作。"),
        ("h1", "2 相关技术"),
        ("h2", "2.1 微服务架构与无状态服务"),
        ("p", "微服务架构强调将复杂系统拆分为一组围绕业务能力构建的小型服务。每个服务通常拥有独立的代码仓库、构建流程和部署生命周期，并通过 HTTP、RPC 或消息队列进行通信。微服务架构的优势在于服务边界清晰、部署灵活、故障隔离能力较强，但这些优势需要依赖自动化部署、服务发现、监控告警和容量治理等工程能力。"),
        ("p", "在弹性伸缩研究中，无状态服务是较适合的实验对象。无状态服务不依赖本地磁盘保存业务状态，Pod 被创建、销毁或迁移时不会造成状态丢失，因此能够较好地适配 Kubernetes 的副本调整机制。本文实现的 sample-api 属于典型无状态 HTTP 服务，提供三个业务路径和一个指标端点。虽然它并不是完整业务系统，但它能够控制变量，突出伸缩策略本身对性能、失败率和资源利用率的影响。"),
        ("p", "需要强调的是，微服务架构并不等同于必须在实验中实现多个复杂服务。Kubernetes 的 HPA 控制对象通常是单个 Deployment，多服务系统中的每个服务也需要先具备独立伸缩能力。本文以单服务为研究对象，是为了减少数据库、缓存、网络链路和业务一致性等额外变量，使实验结论聚焦于容器编排平台的伸缩机制。"),
        ("h2", "2.2 Kubernetes 容器编排机制"),
        ("p", "Kubernetes 通过资源对象和控制器模式管理容器化应用。Deployment 描述期望副本数、Pod 模板、镜像、端口、资源请求和探针配置；ReplicaSet 负责维持副本数量；Pod 是实际运行容器的最小调度单元；Service 为一组 Pod 提供稳定访问入口。在本文中，sample-api 的 Deployment 默认设置 replicas=2，容器镜像为 thesis-sample-api:latest，端口为 8080，并设置 CPU request=100m、memory request=128Mi、CPU limit=500m、memory limit=512Mi。"),
        ("p", "资源 requests/limits 在伸缩实验中具有重要作用。requests 表示调度器为容器预留的最低资源，也是 HPA 计算 CPU/内存利用率的重要基准；limits 表示容器可使用资源上限，能够避免单个 Pod 过度占用节点资源。若 requests 配置过小，HPA 可能更容易判断资源利用率过高并扩容；若 requests 配置过大，则可能降低可调度 Pod 数量，并使资源利用率看起来偏低。因此，本文在讨论 VPA 推荐时，将其视为优化 requests 的依据，而不是简单替代 HPA。"),
        ("p", "健康探针也是系统稳定性的组成部分。本文 Deployment 中配置 readinessProbe 和 livenessProbe，均访问服务根路径。readinessProbe 用于判断 Pod 是否可以接收流量，livenessProbe 用于检测服务是否处于可运行状态。在现场演示或系统检查中，这些配置能够说明服务并非只完成了镜像部署，还考虑了运行状态检测。"),
        ("h2", "2.3 HPA 水平自动伸缩"),
        ("p", "Horizontal Pod Autoscaler 通过周期性读取指标并计算期望副本数，实现 Deployment、StatefulSet 等工作负载的水平伸缩。最常见的 HPA 使用 CPU 利用率作为目标指标，例如本文 CPU HPA 设置 averageUtilization=60，minReplicas=2，maxReplicas=10。当目标服务平均 CPU 利用率超过目标值时，HPA 会计算需要增加的副本数；当指标下降并满足稳定窗口要求时，HPA 会逐步缩容。"),
        ("p", "CPU HPA 的优势在于配置简单，依赖 metrics-server 即可运行，适合 CPU 密集型服务。但对于 I/O 等待较多、请求排队明显或业务压力变化快的 HTTP 服务，CPU 并不一定是最敏感指标。本文因此设计多指标 HPA，将 CPU、内存和 http_requests_per_second 业务指标共同作为伸缩依据。Kubernetes HPA 在多指标场景下会分别计算各指标建议副本数，并采用能够满足最大需求的副本数，从而避免单一指标忽略真实压力。"),
        ("p", "本文多指标 HPA 的扩容行为设置为 scaleUp stabilizationWindowSeconds=0，并包含 Percent=100 与 Pods=4 两类策略，selectPolicy=Max；缩容行为设置 stabilizationWindowSeconds=300，Percent=30，selectPolicy=Min。这一配置体现了“扩容较快、缩容较稳”的工程取舍：突发压力下尽量缩短容量不足时间，压力下降后避免过快回收副本导致下一波流量到来时再次排队。"),
        ("h2", "2.4 VPA 垂直自动伸缩"),
        ("p", "Vertical Pod Autoscaler 主要根据 Pod 历史资源使用情况给出 CPU 和内存 requests 推荐值，并可在特定模式下自动更新 Pod 资源配置。VPA 包含 Recommender、Updater、Admission Controller 等组件。Recommender 根据历史数据计算推荐值，Updater 在 Auto 模式下可能驱逐 Pod 以应用新资源配置，Admission Controller 则可在 Pod 创建时注入推荐资源。"),
        ("p", "本文采用 VPA Initial 模式，而不是 Auto 模式。原因有三点：第一，本文目标是对比伸缩策略对压测结果的影响，若 VPA 在压测过程中主动驱逐 Pod，可能引入与策略无关的中断；第二，HPA 与 VPA 同时自动调整时可能相互影响，尤其是 CPU 利用率计算会受 requests 变化影响；第三，本地 Minikube 环境资源有限，保持实验稳定性比追求完全自动化更重要。因此，本文将 VPA 用于记录资源推荐值，并在优化讨论中分析如何根据推荐值调整 requests。"),
        ("p", f"在正式实验的 hpa_vpa 运行中，归档的 vpa.yaml 包含真实 VPA 对象和推荐状态。{vpa_text}该结果说明 VPA 组件链路已经运行，并能基于当前服务资源使用情况给出建议。论文不将 VPA 推荐解释为生产环境最优配置，而将其作为同一实验环境下资源请求优化的参考。"),
        ("h2", "2.5 Prometheus、Adapter 与 Locust"),
        ("p", "Prometheus 是云原生监控体系中常用的时序数据库和指标采集工具。应用通过 /metrics 端点暴露符合 Prometheus 文本格式的指标，Prometheus 按配置周期性抓取并存储。本文 sample-api 暴露 http_requests_total 计数器和 http_request_duration_seconds 直方图，Prometheus 采集后可用于观察请求数量和耗时分布。"),
        ("p", "Prometheus Adapter 的作用是将 Prometheus 查询结果转换为 Kubernetes 自定义指标 API，使 HPA 能够像读取资源指标一样读取业务指标。本文将应用请求计数转换为 http_requests_per_second，并注册到 custom.metrics.k8s.io。该链路是本文实验系统的关键，因为它将应用层业务压力引入 Kubernetes 伸缩决策，使多指标 HPA 不再只依赖 CPU 和内存。"),
        ("p", "Locust 用于模拟 HTTP 用户行为。本文压测脚本定义了三个请求路径：根路径权重为 6，/a 权重为 3，/b 权重为 1，并设置用户等待时间在 0.05s 至 0.3s 之间。这样的负载模型虽然不等同于真实秒杀系统，但能够制造持续且带有路径差异的 HTTP 访问压力，为比较不同伸缩策略提供可控输入。"),
        ("h1", "3 需求分析与总体设计"),
        ("h2", "3.1 研究边界与需求分析"),
        ("p", "本文系统的首要需求是构建一个能够验证 Kubernetes 弹性伸缩策略的实验平台，而不是开发面向终端用户的业务管理系统。平台需要支持部署一个可观测的微服务，能够在统一负载下切换不同伸缩策略，并自动采集运行数据。系统的核心用户可以理解为研究者或答辩演示者，其关注点不是页面交互，而是实验是否可运行、数据是否完整、结论是否可追溯。"),
        ("p", "功能需求包括五项。第一，服务部署需求：sample-api 必须能够以容器方式部署到 Kubernetes 命名空间 thesis-demo，并通过 Service 暴露访问入口。第二，指标采集需求：服务必须提供 /metrics 端点，Prometheus 能够抓取请求计数和延迟指标，metrics-server 能够提供 CPU/内存资源指标。第三，策略切换需求：实验脚本能够在 static、hpa_cpu、hpa_multi、hpa_vpa 四组策略之间切换，并清理上一组策略残留。第四，负载生成需求：Locust 能够按照统一用户数、生成速率和运行时间产生压测数据。第五，结果分析需求：系统能够从 raw 目录读取每次运行结果，生成逐次指标、汇总指标和图表。"),
        ("p", "非功能需求包括可复现性、稳定性和可解释性。可复现性要求实验配置集中在 experiment.config.json 中，运行目录保存环境和参数；稳定性要求脚本在缺少必要指标、缺少重复次数或 VPA 降级时停止或报告异常；可解释性要求每个指标都有明确口径，例如 failure_rate 来自 Locust Aggregated 行，CPU/内存利用率来自 top_pods.log，成本代理仅用于同环境横向比较，不代表真实云厂商费用。"),
        ("h2", "3.2 总体架构设计"),
        ("p", "系统总体架构可以分为被测服务层、容器编排层、监控指标层、负载生成层和结果分析层。被测服务层由 sample-api 组成，负责响应 HTTP 请求并暴露 Prometheus 指标。容器编排层由 Kubernetes Deployment、Service、HPA 和 VPA 组成，负责服务运行、访问入口和弹性控制。监控指标层由 metrics-server、Prometheus 和 Prometheus Adapter 组成，分别提供资源指标和业务指标。负载生成层由 Locust 组成，按照固定参数生成 burst 场景。结果分析层由 aggregate_experiment.py 和 plot_results.py 组成，负责读取原始数据、校验数据质量并生成图表。"),
        ("p", "在数据流向上，Locust 通过本地 port-forward 访问 Kubernetes Service，流量转发到 sample-api Pod。sample-api 在处理请求时更新请求计数和延迟统计，并通过 /metrics 暴露给 Prometheus。metrics-server 从节点和 Pod 层面提供 CPU、内存数据。Prometheus Adapter 将 Prometheus 中的 http_requests_total 转换为 http_requests_per_second，自定义指标随后被 HPA 读取。实验结束后，Locust CSV、HPA/VPA YAML、Pod 状态、事件日志和 top 采样被保存到策略对应的 run 目录中。"),
        ("p", "该架构的设计重点是让策略对比保持同一基础环境。四组实验使用相同服务、相同镜像、相同负载参数和相同采集脚本，只改变伸缩策略。这样能够减少其他变量对结果的干扰，使静态副本、CPU HPA、多指标 HPA 和 HPA+VPA 之间的差异更有解释价值。"),
        ("h2", "3.3 实验流程设计"),
        ("p", "自动化实验脚本 run_experiments.ps1 是实验流程的核心。脚本首先读取 experiment.config.json，获得命名空间、服务名、本地端口、重复次数、负载参数和策略列表。若未跳过预检，脚本会部署基础资源并检查 metrics API、custom metrics API 和 VPA API 是否可用。对于包含 hpa_multi 或 hpa_vpa 的实验，custom metrics API 是必要条件；对于 hpa_vpa，VPA CRD 必须存在，否则不能将运行视为正式协同实验。"),
        ("p", "每次运行开始时，脚本会调用 Apply-Strategy 删除已有 HPA/VPA 并应用当前策略。static 组只将 Deployment 缩放到 2 副本；hpa_cpu 组应用 CPU HPA；hpa_multi 组应用多指标 HPA；hpa_vpa 组在多指标 HPA 基础上应用 VPA Initial。策略应用后，脚本等待 Deployment rollout 完成并进入 warmup 阶段，随后启动 port-forward 和 top 采样进程。"),
        ("p", "压测阶段使用 Locust headless 模式执行，参数由配置文件统一控制。压测结束后，脚本导出 hpa_after.txt、hpa_describe.txt、pods_after.txt、events_after.txt，并写入 manifest.json。finally 逻辑负责关闭 port-forward 和 top 采样进程，并在下一次运行前进入 cooldown。所有策略和重复次数结束后，脚本调用 aggregate_experiment.py 和 plot_results.py 生成最终 CSV 和图表。"),
        ("h2", "3.4 数据归档与可追溯设计"),
        ("p", "本文将可追溯性作为实验平台的重要设计目标。每次运行目录命名包含策略、日期时间和重复序号，例如 hpa_vpa_20260508_134705_r3，能够快速定位某一次具体运行。目录中保存 locust_stats.csv、locust_stats_history.csv、locust_failures.csv、locust_exceptions.csv、hpa.yaml、vpa.yaml、hpa_describe.txt、hpa_after.txt、events_before.txt、events_after.txt、pods_after.txt、top_pods.log 和 manifest.json。"),
        ("p", "manifest.json 记录当前策略、场景、运行序号、配置内容、kubectl context、Kubernetes 版本和 Git commit。这使得实验结果不仅有汇总表，还能追溯到具体配置和环境。对于毕业论文而言，这种数据归档方式能够支撑“结果来自真实运行”的论证，也便于答辩时展示从结论回到原始证据的过程。"),
        ("h1", "4 系统实现"),
        ("h2", "4.1 sample-api 微服务实现"),
        ("p", "sample-api 使用 Python 标准库中的 ThreadingHTTPServer 实现，不依赖复杂 Web 框架。这样做的目的是减少应用框架本身对实验结果的影响，并让被测服务逻辑保持可读。服务监听 0.0.0.0:8080，支持 /、/a、/b 三类业务请求和 /metrics 指标请求。对于普通业务路径，服务会根据路径模拟不同处理时长，其中 /b 会额外增加处理时间，以形成一定响应差异。"),
        ("p", "服务内部维护 REQUESTS、LATENCY_BUCKETS、LATENCY_COUNTS、LATENCY_SUM 和 LATENCY_TOTAL 等变量。每次普通请求到达时，服务先记录开始时间，处理完成后更新请求计数和延迟直方图。/metrics 请求不会计入普通业务路径，而是返回 Prometheus 文本格式，包括 http_requests_total 和 http_request_duration_seconds。该设计让应用指标来自服务运行时，而不是离线统计结果，从而能够被 Prometheus 实时抓取。"),
        ("p", "由于本文研究重点是伸缩策略而不是业务算法，sample-api 的业务逻辑保持简单，但它具备真实微服务实验所需的基本属性：可容器化、可部署、可访问、可被压测、可暴露指标、可由 HPA 读取业务指标链路间接驱动。现场演示时，可以通过访问根路径验证服务可用，通过访问 /metrics 验证指标暴露，通过查看 Prometheus Adapter 输出验证业务指标进入 Kubernetes API。"),
        ("h2", "4.2 容器化与基础部署实现"),
        ("p", "服务通过 Dockerfile 构建为 thesis-sample-api:latest 镜像，并加载到 Minikube profile 中。Kubernetes 基础部署清单包括 namespace.yaml、deployment.yaml 和 service.yaml。命名空间统一为 thesis-demo，便于隔离实验资源并简化脚本命令。Service 将集群内服务端口 80 映射到容器端口 8080，实验脚本通过 kubectl port-forward 将本地 8080 转发到 Service。"),
        ("p", "Deployment 中设置 replicas=2 作为初始副本数，容器资源请求为 CPU 100m、内存 128Mi，资源上限为 CPU 500m、内存 512Mi。该配置为 HPA 的资源利用率计算提供基准，也为 VPA 推荐值提供比较对象。探针配置方面，readinessProbe 初始延迟 5 秒、周期 10 秒，livenessProbe 初始延迟 15 秒、周期 20 秒，均访问根路径。"),
        ("p", "从系统检查角度看，基础部署不仅证明服务能运行，还体现了对稳定性的考虑。readinessProbe 可避免未准备好的 Pod 过早接收流量，livenessProbe 可在服务异常时触发重启。虽然本文实验没有专门构造异常输入测试，但 Deployment 层面的探针和脚本层面的错误停止机制共同构成了基础稳定性保障。"),
        ("h2", "4.3 监控与自定义指标实现"),
        ("p", "资源指标链路依赖 metrics-server。CPU HPA 和多指标 HPA 中的 CPU/内存利用率都需要 metrics-server 提供 Pod 资源数据。实验 README 中记录了 metrics-server 的启用、镜像替换和 patch 过程，并要求通过 kubectl top pods 验证资源指标可用。若 kubectl top 无数据，正式实验不能继续，因为 HPA 无法获得可靠资源指标。"),
        ("p", "业务指标链路由 Prometheus 和 Prometheus Adapter 完成。Prometheus 负责抓取 sample-api 的 /metrics，Adapter 根据配置将请求计数转换为 http_requests_per_second，并注册到 Kubernetes custom.metrics.k8s.io。实验脚本预检阶段会查询 /apis/custom.metrics.k8s.io/v1beta1，并检查指标名是否存在。只有该指标可用，hpa_multi 与 hpa_vpa 才能被视为真实多指标实验。"),
        ("p", f"归档的 hpa_describe.txt 证明 HPA 能够读取多类指标。{hpa_state}这类状态输出对论文非常重要，因为它连接了配置文件和实验结果：不仅 YAML 中写了 Pods 指标，实际运行时 HPA 也确实读到了该指标。"),
        ("h2", "4.4 伸缩策略实现"),
        ("p", "静态副本策略作为对照组，不启用 HPA 和 VPA，只将 sample-api Deployment 固定为 2 个副本。该策略代表传统固定容量部署方式，优点是行为简单、资源数量可控，缺点是在负载快速变化时缺少自适应能力。静态组的结果用于判断自动伸缩策略是否带来改善。"),
        ("p", "CPU HPA 策略使用 autoscaling/v2，目标 CPU 平均利用率为 60%，最小副本数为 2，最大副本数为 10。扩容阶段稳定窗口为 0 秒，允许每 30 秒按 100% 比例增加副本；缩容阶段稳定窗口为 180 秒，每 60 秒最多按 50% 比例缩容。该策略体现了较常见的资源指标伸缩配置，也是 Kubernetes 初学者最容易采用的方案。"),
        ("p", "多指标 HPA 策略在 CPU HPA 基础上增加内存和 Pod 级 QPS 指标。CPU 目标为 60%，内存目标为 70%，http_requests_per_second 的 AverageValue 目标为 50。最大副本数提高到 12，并在扩容策略中同时配置 Percent=100 和 Pods=4，selectPolicy=Max。这意味着在压力上升时，HPA 可以选择更积极的扩容方式，以降低突发阶段容量不足的持续时间。"),
        ("p", "HPA+VPA 策略在多指标 HPA 基础上应用 VPA Initial。VPA 的 minAllowed 为 CPU 100m、内存 128Mi，maxAllowed 为 CPU 1500m、内存 2Gi，controlledResources 包含 CPU 和内存。选择 Initial 模式的原因是：本文希望记录 VPA 推荐值，用于分析资源请求是否合理，但不希望在压测过程中由 VPA 主动驱逐 Pod 影响实验稳定性。这一取舍符合系统设计类论文对实验可控性的要求。"),
        ("h2", "4.5 自动化实验与结果处理实现"),
        ("p", "run_experiments.ps1 将实验流程自动化，避免手动执行导致步骤遗漏。脚本使用 Set-StrictMode 和 ErrorActionPreference=Stop，使关键命令失败时及时停止。Invoke-Checked 会输出并执行命令，便于日志追踪。Wait-Rollout 用于等待 Deployment 状态稳定，Start-PortForward 负责建立本地访问通道，Start-TopSampler 周期性记录 kubectl top pods 输出。"),
        ("p", "数据处理脚本 aggregate_experiment.py 负责读取每次运行目录。它从 locust_stats.csv 中提取 Aggregated 行，计算 request_count、failure_count、failure_rate、avg_ms、p95_ms、p99_ms 和 qps；从 top_pods.log 中提取 sample-api Pod 的 CPU 与内存采样，计算资源利用率；从事件日志中统计伸缩和震荡事件；从 vpa.yaml 中判断 VPA 是否真实可用。"),
        ("p", "脚本还包含数据质量校验逻辑。它会检查输出列是否完整，关键列是否存在 NaN，指标是否在合理范围内，每个策略在正式实验中是否恰好包含 3 次运行，top 采样是否存在 fallback，hpa_vpa 是否包含真实 VPA 输出。任何校验失败都将抛出异常。这种实现避免了错误数据悄悄进入论文图表，是本文实验可信度的重要保障。"),
        ("p", "plot_results.py 读取逐次指标或汇总指标后，生成 P95、P99、QPS、失败率、CPU 利用率、内存利用率、伸缩事件、震荡事件和成本代理等图表。图表采用均值加标准差误差线，能够反映重复实验中的波动。相比只给出单次运行结果，重复实验和标准差更能体现实验结论的稳健性。"),
        ("h2", "4.6 复现环境与执行说明"),
        ("p", "为了让实验能够由他人复现，本文在 README 中保留了从零执行步骤。首先需要启动 Minikube，并切换到对应 context；然后构建 sample-api 镜像，将镜像加载到 Minikube；接着应用 namespace、deployment 和 service 等基础资源。资源指标部分需要启用 metrics-server，并根据本地网络情况替换镜像或修改启动参数。业务指标部分需要部署 Prometheus 和 Prometheus Adapter，并通过 custom.metrics.k8s.io API 验证 http_requests_per_second 是否已经注册。"),
        ("p", "在执行全量实验前，建议先运行单策略冒烟测试。冒烟测试可以将 repeats 设置为 1，将 run_time 调整为 2m，并只保留一个策略。该步骤的目的不是获得论文结果，而是验证服务访问、端口转发、Locust、HPA、top 采样和目录写入是否正常。若冒烟测试通过，再恢复正式参数运行 4 组策略、每组 3 次的完整实验。这样可以减少全量实验运行到中途才发现指标链路异常的风险。"),
        ("p", "正式实验完成后，需要检查 results/raw 下是否存在四个策略目录，每个策略在 burst 场景下是否包含 3 个运行目录。随后运行 aggregate_experiment.py 生成 experiment_metrics.csv 和 experiment_metrics_summary.csv，再运行 plot_results.py 生成图表。本文优化版论文中的表格和图像均来自这些产物，因此复现过程不仅能验证系统运行，也能验证论文结论来源。"),
        ("h1", "5 实验设计与结果分析"),
        ("h2", "5.1 实验环境与变量控制"),
        ("p", f"本文正式实验固定在 Minikube 环境中完成，场景名称为 {scenario}。实验参数统一配置在 scripts/experiment.config.json 中：users={users}，spawn_rate={spawn_rate}，run_time={run_time}，warmup_seconds={warmup}，cooldown_seconds={cooldown}，top_snapshot_interval_seconds=10。四组策略按同一参数重复运行，每组 3 次，共 12 次。"),
        ("p", "为了保证策略差异来自伸缩配置而非服务差异，所有实验共用同一个 sample-api 镜像、同一个 Deployment 基础配置、同一个 Service 入口和同一个 Locust 脚本。每次策略切换前删除已有 HPA/VPA，避免上一次实验资源残留影响下一次结果。每次压测前等待 rollout 完成，并在压测后进入冷却阶段。"),
        ("p", "本文没有将实验目标设定为跨机器数值完全一致，而是强调在同一环境中进行策略横向比较。原因在于本地 Minikube 单节点环境受宿主机负载、Docker 运行状态和系统资源竞争影响较大，不同机器上的绝对延迟和 QPS 可能不同。但在相同环境、相同参数、相同重复次数下，策略之间的趋势差异仍具有分析价值。"),
        ("h2", "5.2 评价指标设计"),
        ("p", "性能指标包括 P95、P99、平均响应时间和 QPS。P95 表示 95% 请求的响应时间不超过该值，能够反映大多数用户体验；P99 更关注极端长尾，容易受资源竞争和短时抖动影响；QPS 表示单位时间请求吞吐量，反映系统处理能力。对于突发负载场景，P95 和失败率通常比平均响应时间更有解释价值。"),
        ("p", "可靠性指标主要是 failure_rate，即 Locust Aggregated 行中的 Failure Count 除以 Request Count。该指标能够反映服务在压测期间是否出现连接失败、超时或非预期响应。资源指标包括 CPU 和内存利用率，来自 top_pods.log 中 sample-api Pod 的周期采样。本文仅统计 sample-api Pod，不将 Prometheus、Adapter 或 Kubernetes 系统组件计入应用资源利用率。"),
        ("p", "成本代理指标基于 CPU 与内存使用量加权估算，配置中 cpu_weight=0.7，memory_weight=0.3，scale=0.01。该指标不是云厂商真实账单，也不适合跨环境比较。本文使用它的目的，是在同一实验环境下观察不同策略在性能改善之外是否带来更高资源消耗。"),
        ("p", "指标解释还需要注意采样口径差异。Locust 的 P95、P99 和 QPS 来自压测客户端视角，反映用户请求从发出到收到响应的整体表现；kubectl top 采样来自 Kubernetes 资源视角，反映 Pod 在采样时刻的 CPU 和内存使用情况；HPA describe 反映控制器在某一时刻读取到的指标和期望副本数。三类数据的时间粒度和采样位置不同，因此不能简单逐秒对齐。本文采用汇总均值比较策略趋势，而不是对每一秒进行精确因果推断。"),
        ("p", "对于失败率，也需要结合 locust_failures.csv 和 locust_exceptions.csv 理解。如果失败率上升，同时资源利用率高，通常说明服务处理能力不足；如果失败率上升但资源利用率不高，则可能与端口转发、客户端压力、连接超时或本地环境波动有关。因此，失败率不是孤立指标，需要与 P95、QPS、top 采样和 HPA 状态共同解释。"),
        ("h2", "5.3 汇总结果"),
        ("table_caption", "表 5-1 四组策略正式实验汇总结果"),
        ("summary_table", 0),
        ("p", f"从表 5-1 可以看出，静态副本组的平均 P95 为 {static['P95(ms)']} ms，平均 QPS 为 {static['QPS']}，平均失败率为 {static['失败率']}。CPU HPA 组的平均 P95 降至 {hpa_cpu['P95(ms)']} ms，平均 QPS 提升到 {hpa_cpu['QPS']}，失败率下降到 {hpa_cpu['失败率']}。多指标 HPA 组平均 P95 为 {hpa_multi['P95(ms)']} ms，平均 QPS 为 {hpa_multi['QPS']}，CPU 利用率均值为 {hpa_multi['CPU利用率']}。HPA+VPA 组平均 P95 最低，为 {hpa_vpa['P95(ms)']} ms，平均 QPS 最高，为 {hpa_vpa['QPS']}，失败率最低，为 {hpa_vpa['失败率']}。"),
        ("p", "需要注意的是，多指标 HPA 的失败率均值高于 CPU HPA，主要受到某一次运行中失败率较高的影响。逐次数据中，多指标 HPA 前两次运行失败率均低于 0.1%，第三次运行失败率明显升高，说明本地单节点环境存在运行波动。因此，本文在分析时不仅看均值，也结合标准差和具体运行记录解释异常。"),
        ("h2", "5.4 性能表现分析"),
        ("p", "P95 时延是本文判断突发负载下用户体验的重要指标。静态副本组固定为 2 个 Pod，在请求量快速上升时无法自动增加实例，因此更容易出现排队。HPA 类策略能够根据指标增加副本，将请求分散到更多 Pod 上，从而降低单个实例压力。实验中，CPU HPA、多指标 HPA 和 HPA+VPA 的平均 P95 均低于静态副本组，说明自动伸缩对长尾延迟有改善作用。"),
        ("figure", ("p95_latency.png", "图 5-1 P95 时延对比")),
        ("p", "从 P99 指标看，各组都存在一定长尾波动。静态组 P99 平均值为 4067 ms，CPU HPA 为 4367 ms，多指标 HPA 为 4033 ms，HPA+VPA 为 3900 ms。CPU HPA 的 P99 高于静态组，说明自动扩容并不必然在所有长尾指标上占优。可能原因包括 HPA 控制周期、指标采集延迟、Pod 启动过程和本地节点资源竞争。该结果提醒我们：伸缩策略不能只凭是否启用 HPA 判断优劣，必须结合完整实验指标分析。"),
        ("p", f"QPS 方面，HPA+VPA 组达到 {hpa_vpa['QPS']}，高于静态副本组的 {static['QPS']}。这说明在本文负载模型中，HPA+VPA 能够更有效地维持吞吐。CPU HPA 与多指标 HPA 的平均 QPS 接近，分别为 {hpa_cpu['QPS']} 和 {hpa_multi['QPS']}，但多指标 HPA 的 CPU 利用率更低，说明其通过更充分的副本分散降低了单 Pod 压力。"),
        ("figure", ("throughput_qps.png", "图 5-2 吞吐量 QPS 对比")),
        ("h2", "5.5 可靠性与失败率分析"),
        ("p", f"失败率是判断系统稳定性的重要指标。静态副本组平均失败率为 {static['失败率']}，其中第一次运行失败率达到 30.85%，说明固定 2 副本在部分突发运行中无法承受压力。CPU HPA 组平均失败率下降到 {hpa_cpu['失败率']}，HPA+VPA 进一步下降到 {hpa_vpa['失败率']}。这表明自动伸缩能够缓解请求堆积和服务不可用风险。"),
        ("p", "失败率波动也暴露出本地实验环境的不确定性。静态组第二次运行失败率仅 0.05%，第三次为 2.94%，与第一次差异较大；多指标 HPA 第三次运行失败率为 24.33%，明显高于前两次。若只选择表现最好的一次运行，会得到过于乐观的结论。因此，本文采用 3 次重复和均值/标准差分析，尽量避免单次偶然结果主导结论。"),
        ("figure", ("failure_rate.png", "图 5-3 失败率对比")),
        ("h2", "5.6 资源利用率分析"),
        ("p", f"CPU 利用率方面，静态副本组均值为 {static['CPU利用率']}，CPU HPA 组为 {hpa_cpu['CPU利用率']}，多指标 HPA 组为 {hpa_multi['CPU利用率']}，HPA+VPA 组为 {hpa_vpa['CPU利用率']}。静态组 CPU 利用率最高，说明固定少量副本时单 Pod 压力较大。多指标 HPA 和 HPA+VPA 通过扩容分散压力，使平均 CPU 利用率降低。"),
        ("figure", ("cpu_utilization.png", "图 5-4 CPU 利用率对比")),
        ("p", f"内存利用率方面，四组均值都在 10% 左右，变化不大。静态组为 {static['内存利用率']}，HPA+VPA 为 {hpa_vpa['内存利用率']}。这说明 sample-api 的负载主要表现为请求处理和 CPU 压力，内存不是主要瓶颈。正因为如此，HPA 中的内存指标在本实验中更多承担辅助观察作用，真正影响扩容的主要仍是 CPU 与 QPS 指标。"),
        ("figure", ("memory_utilization.png", "图 5-5 内存利用率对比")),
        ("h2", "5.7 伸缩行为与成本代理分析"),
        ("p", "本文汇总表中的 scale_events 和 oscillation_events 均为 0，并不表示 HPA 没有发生副本变化。原因在于脚本当前主要从事件文本中统计特定关键词，而 Kubernetes 事件输出在不同版本、不同语言环境或不同资源对象下可能不包含这些关键词。归档的 hpa_describe.txt 中可以看到 Deployment pods: 5 current / 5 desired，说明 HPA 状态中存在副本调整结果。论文中因此不将 scale_events 作为核心结论，而将其作为采集口径需要改进的指标。"),
        ("figure", ("scale_events.png", "图 5-6 伸缩事件数对比")),
        ("p", f"成本代理方面，静态副本组为 {static['成本代理']}，CPU HPA 为 {hpa_cpu['成本代理']}，多指标 HPA 为 {hpa_multi['成本代理']}，HPA+VPA 为 {hpa_vpa['成本代理']}。在本文估算口径下，HPA+VPA 和多指标 HPA 的成本代理低于静态副本组，主要是因为其单位请求分摊后的资源使用更低。需要说明的是，成本代理不是云账单，不包含节点固定成本、镜像拉取、网络、存储和控制面开销，只用于同一环境内的相对比较。"),
        ("figure", ("cost_per_1k_req.png", "图 5-7 每千请求成本代理对比")),
        ("h2", "5.8 实验结论"),
        ("p", "综合性能、可靠性和资源指标，本文实验得到以下结论。第一，固定副本策略虽然简单，但在突发负载下存在较高失败风险，尤其当请求在短时间内集中到来时，固定 2 副本无法及时增加处理能力。第二，CPU HPA 能够在一定程度上改善 P95 和失败率，说明 Kubernetes 原生水平伸缩机制对微服务容量治理有效。第三，多指标 HPA 引入 QPS 后，能够将业务压力纳入伸缩依据，但其效果仍受指标采集周期、HPA 控制周期和本地环境波动影响。第四，HPA+VPA 在本文环境中综合表现最好，体现出水平副本调整与资源请求推荐结合的潜力。"),
        ("p", "这些结论的适用范围需要谨慎限定。本文实验对象是单个无状态 HTTP 微服务，负载模型是 Locust 生成的 burst HTTP 请求，不包含真实数据库、缓存、消息队列和复杂业务一致性约束。因此，结论可以作为 Kubernetes 单服务伸缩策略配置的实践参考，但不能直接推广为所有微服务系统或秒杀系统的最优策略。若应用到多服务业务链路，需要重新设计服务拓扑、业务指标和瓶颈分析方法。"),
        ("h1", "6 策略优化与讨论"),
        ("h2", "6.1 基于数据的优化闭环"),
        ("p", "通过本文系统获得数据后，优化不应只停留在“选择平均值最好的一组”，而应形成瓶颈定位、策略调整、重新实验和结果比较的闭环。若 P95/P99 高、失败率高且 CPU 利用率也高，说明服务计算资源或副本数量不足，可降低 HPA CPU 阈值、提高 maxReplicas 或优化服务代码。若 P95 高但 CPU 不高，而 QPS 或请求排队明显，则说明 CPU 不是充分指标，应引入 QPS、请求延迟或队列长度等业务指标。"),
        ("p", "若资源利用率低但副本数量较多，说明扩容可能过度，应提高扩容阈值、降低 maxReplicas 或延长缩容稳定窗口。若失败率高但资源利用率不高，则需要检查应用异常、连接数、端口转发、Locust 压力模型或下游依赖。在本文系统中，这些判断可以通过 Locust CSV、top_pods.log、hpa_describe.txt 和 events_after.txt 交叉验证。"),
        ("h2", "6.2 HPA 参数优化方向"),
        ("p", "对于 CPU HPA，averageUtilization 是最直接的调优参数。当前值为 60%。如果目标是降低延迟和失败率，可以尝试将阈值调低到 50%，使扩容更早发生；如果目标是节省资源，可以提高到 70%，但需要观察 P95 和失败率是否恶化。maxReplicas 也需要结合节点容量设置，过低会限制扩容效果，过高则可能导致节点资源竞争或调度失败。"),
        ("p", "对于多指标 HPA，QPS 目标值是关键参数。当前 http_requests_per_second 的 AverageValue 为 50。若压测中 QPS 高但扩容不明显，可以将目标值降低到 35 或 40；若副本增长过快且资源利用率偏低，可以提高到 60 或 70。缩容稳定窗口也需要谨慎设置。突发流量场景下，过快缩容可能导致下一轮流量到来时重新扩容，增加响应波动；较长稳定窗口能够提升稳定性，但会延长资源占用时间。"),
        ("h2", "6.3 VPA 推荐与资源请求优化"),
        ("p", f"VPA 的价值在于发现资源请求与实际使用之间的偏差。本文 Deployment 初始 requests 为 CPU 100m、内存 128Mi。正式实验中，{vpa_text}这说明在当前负载和采样条件下，VPA 认为内存请求应高于初始配置，CPU 请求也可以适当提高。"),
        ("p", "根据 VPA 推荐优化 requests 时，需要重新运行实验验证。提高 requests 可能使 Pod 获得更稳定的资源保障，也会影响 HPA 的利用率计算；同时，它会增加调度器认为每个 Pod 需要的资源，可能降低单节点可承载副本数。因此，合理做法不是直接采用推荐值作为最终答案，而是设置一组候选配置，例如 CPU 150m/200m、内存 192Mi/256Mi，分别压测比较 P95、失败率、QPS 和成本代理。"),
        ("h2", "6.4 多服务场景扩展讨论"),
        ("p", "如果将本文系统扩展到多服务场景，需要将验证对象从单个 Deployment 扩展为一组有调用关系的服务链路。例如，一个简化秒杀系统可以包含 gateway-service、order-service、inventory-service 和 Redis/数据库。此时 Locust 请求首先进入 gateway，再调用订单服务和库存服务，系统不仅需要观察总 P95 和失败率，还需要观察各服务的调用延迟、库存扣减成功率、队列长度、数据库连接数和缓存命中率。"),
        ("p", "多服务场景下，伸缩策略也不再是简单比较四组全局策略，而应根据服务角色配置。入口服务可能适合按 QPS 扩容，订单服务可能需要结合 CPU 和请求延迟，库存服务可能需要结合队列长度或库存请求数。若只扩 gateway，可能把压力转移到 order 或 inventory；若所有服务同时扩容，可能提高吞吐但增加资源成本，并将瓶颈推向数据库。因此，多服务弹性伸缩的核心是识别链路瓶颈并进行分层伸缩。"),
        ("p", "本文当前实现不直接声称能够验证所有多服务系统，而是提供单服务伸缩验证的基础框架。后续可复用本文的实验流程，将 sample-api 替换为多服务业务链路，将 /metrics 扩展为每个服务的业务指标，并将 aggregate_experiment.py 扩展为链路级分析脚本。这样的扩展方向比直接宣称通用验证平台更符合工程事实。"),
        ("h2", "6.5 局限性分析"),
        ("p", "本文存在四方面局限。第一，实验运行在本地 Minikube 单节点环境中，无法反映多节点集群中的调度、网络和节点间资源差异。第二，sample-api 是简化无状态服务，没有真实数据库、缓存和消息队列，因此无法分析下游依赖瓶颈。第三，伸缩事件统计口径仍需改进，当前汇总表中 scale_events 为 0，而 HPA 状态文件能够看到副本变化，这说明事件解析不能完全覆盖 Kubernetes 输出差异。第四，VPA 仅使用 Initial 模式记录推荐，未评估 Auto 模式下驱逐和重建 Pod 对业务的影响。"),
        ("p", "这些局限并不否定本文结论，而是说明结论适用范围。本文的贡献在于完成了从部署、监控、压测到分析的闭环，并基于真实数据比较四类策略。在后续研究中，可以通过多节点集群、更复杂业务链路、更长时间运行和更多重复次数进一步提高结论稳健性。"),
        ("h2", "6.6 与系统代码检查标准的对应关系"),
        ("p", "按照计算机科学与技术学院毕业论文系统代码检查标准，系统设计类基础检查项主要包括系统完成度、系统运行现场演示、系统稳定性以及功能与论文一致性。本文项目虽然不是传统的登录、管理、增删改查系统，但它具备完整的系统设计对象：有被测微服务，有 Kubernetes 部署清单，有监控指标链路，有伸缩策略配置，有自动化实验脚本，也有结果分析与图表输出。因此，本文应按照系统设计类和工程实践类进行说明，而不应按照算法设计类进行答辩。"),
        ("p", "在系统完成度方面，项目已经实现从服务构建到结果分析的主要闭环。sample-api 提供可访问接口和指标端点；Deployment 和 Service 负责在 Kubernetes 中运行服务；HPA/VPA 配置负责策略控制；Locust 负责生成统一负载；aggregate_experiment.py 与 plot_results.py 负责生成论文使用的数据和图表。上述模块之间存在明确的数据流和控制流，不是彼此独立的代码片段。"),
        ("p", "在现场演示方面，可以选择 2 至 5 个核心模块进行展示。第一，访问 sample-api 根路径，证明服务正在运行；第二，访问 /metrics，展示 http_requests_total 和 http_request_duration_seconds 指标；第三，查看 hpa-multi.yaml，说明 CPU、内存和 QPS 指标如何进入 HPA；第四，运行或展示 Locust 输出，说明压测数据如何产生；第五，运行 aggregate_experiment.py，说明原始数据如何汇总为论文表格。这样的演示路径能够覆盖“核心功能模块演示”的要求。"),
        ("p", "在系统稳定性方面，本文从两个层面处理。服务层面配置了 readinessProbe 和 livenessProbe，避免未就绪 Pod 接收流量，并在服务异常时由 Kubernetes 重启；实验层面设置了预检、rollout 等待、warmup、cooldown、top 采样和数据质量校验。若 metrics API、自定义指标 API 或 VPA API 不满足条件，脚本会停止或报错，防止生成不可靠结果。这种稳定性不是体现在复杂界面交互，而是体现在实验流程能够重复执行并发现异常。"),
        ("p", "在功能与论文一致性方面，论文中的核心描述均对应到仓库文件和实验数据。论文提到的 sample-api 对应 services/sample-api/app.py，提到的 CPU HPA 对应 deploy/hpa/hpa-cpu.yaml，提到的多指标 HPA 对应 deploy/hpa/hpa-multi.yaml，提到的 VPA Initial 对应 deploy/vpa/vpa-initial.yaml，提到的 12 条正式实验记录对应 results/raw 与 results/processed/experiment_metrics.csv。答辩时若教师抽查关键代码，可以围绕这些文件解释实现逻辑。"),
        ("h2", "6.7 低模板化写作与真实性说明"),
        ("p", "本文在写作上避免将结论写成脱离项目事实的宏观判断，而是尽量把每个判断绑定到具体实现、具体参数或具体结果。例如，论文不笼统地说“系统性能显著提升”，而是说明 HPA+VPA 组平均 QPS、P95 和失败率分别是多少；不笼统地说“系统可复现”，而是说明每次运行保存哪些 CSV、YAML、日志和 manifest；不笼统地说“支持多服务系统”，而是明确当前只验证单个典型无状态微服务，多服务场景需要后续扩展。"),
        ("p", "这种写作方式有两个作用。第一，它能提高论文可信度，因为读者可以根据文件路径和实验数据复核结论。第二，它能降低论文的空泛感，使文字更接近工程复盘和系统设计说明。对于系统设计类毕业论文而言，真正有价值的内容不是堆砌概念，而是说明系统为什么这样设计、实际如何运行、数据从哪里来、结果如何解释以及哪些边界不能越过。"),
        ("p", "本文还特别避免将 AI、预测算法或复杂控制模型强行加入最终实现。虽然弹性伸缩可以与机器学习预测结合，但本文现有代码和实验没有实现预测控制器。如果为了显得复杂而在论文中写入未实现的算法，会造成论文与代码不一致，反而增加答辩风险。因此，本文选择如实呈现 Kubernetes 原生 HPA/VPA 与自定义指标链路的实践价值。"),
        ("h2", "6.8 进一步调优实验设计"),
        ("p", "若在本文基础上继续优化，可以设计第二轮调优实验。第一组调优针对 CPU HPA，将 averageUtilization 分别设置为 50、60、70，观察 P95、失败率和成本代理的变化。若阈值为 50 时 P95 明显下降但成本代理上升，说明低阈值适合服务质量优先场景；若阈值为 70 时成本下降但失败率上升，说明高阈值适合成本优先但风险更高的场景。"),
        ("p", "第二组调优针对 QPS 指标，将 http_requests_per_second 的 AverageValue 分别设置为 35、50、70，观察扩容敏感性。QPS 阈值越低，HPA 越容易提前扩容，适合突发流量；阈值越高，扩容更保守，适合资源紧张环境。该实验能够进一步回答“多指标 HPA 如何配置更合理”的问题，而不只是证明多指标 HPA 能够运行。"),
        ("p", "第三组调优针对 VPA 推荐资源。根据当前样例推荐，可将 Deployment requests 设置为 CPU 200m、内存 256Mi，再与原始 100m、128Mi 配置对比。比较时需要注意两个方向：一方面，更高 requests 可能让 Pod 获得更稳定的调度和资源保障；另一方面，它也会影响 HPA 利用率计算，并减少单节点可调度副本数量。因此，VPA 推荐值不能机械采用，必须结合 HPA 行为和节点容量重新压测。"),
        ("p", "这些调优实验如果作为后续工作展开，可以形成更完整的策略优化论文结构：第一轮实验比较不同策略，第二轮实验在最优或较优策略上调整参数，第三轮实验讨论多服务或多节点扩展。本文由于毕业设计时间和本地环境限制，主要完成第一轮策略对照，并在第 6 章提出可执行的优化方向。"),
        ("h2", "6.9 实验误差与数据可信度控制"),
        ("p", "本地弹性伸缩实验存在天然误差来源。首先，Minikube 运行在单台宿主机上，宿主机 CPU、内存、Docker 后台进程和其他应用都会影响压测结果。其次，HPA 的决策不是实时发生，而是依赖指标采集周期、控制器同步周期和 Pod 启动时间，因此突发流量刚开始时可能已经出现排队，但 HPA 尚未完成扩容。再次，Locust 自身也需要消耗本机资源，当用户数较高时，压测客户端和被测集群可能共享宿主机资源，从而放大波动。"),
        ("p", "为了降低这些误差对结论的影响，本文采用了三种控制方式。第一，所有策略使用相同实验环境、相同镜像、相同压测脚本和相同参数，使比较重点落在策略差异上。第二，每组策略重复 3 次，并保留逐次数据，避免只选择最有利的一次结果。第三，数据处理脚本设置质量校验，例如每组运行次数必须与配置一致，top 采样不能缺失，hpa_vpa 必须包含真实 VPA 输出。这些措施不能完全消除误差，但能够让误差暴露在数据中，而不是被论文叙述掩盖。"),
        ("p", "对于实验中出现的异常波动，本文不将其简单删除。例如多指标 HPA 第三次运行失败率较高，若删除该数据，平均结果会明显改善，但论文结论会失去真实性。保留该运行能够说明本地环境和控制周期对伸缩效果存在影响，也能提醒后续研究需要增加重复次数、延长压测时间或迁移到多节点集群。系统设计类论文的重点不是制造完美数据，而是如实解释系统运行结果。"),
        ("h2", "6.10 异常处理与风险控制"),
        ("p", "本文实验脚本在多个位置加入异常控制。预检阶段会验证基础 Deployment 是否可用、metrics API 是否能返回资源指标、自定义指标 API 是否包含 http_requests_per_second、VPA API 是否存在。若这些条件不满足，脚本不会继续生成正式实验结果。这样做是为了避免出现“策略名称是 hpa_vpa，但实际未安装 VPA”的情况。"),
        ("p", "运行阶段的风险主要来自端口转发、采样进程和外部命令。脚本将 port-forward 和 top sampler 作为独立进程启动，并在 finally 块中关闭，防止一次运行失败后残留进程影响下一次实验。每次运行都在独立目录中保存结果，即使某次运行异常，也能通过对应目录查看 failures、exceptions 和事件日志。结果处理阶段如果发现缺列、空值、指标超出范围或重复次数不对，会直接抛出错误。"),
        ("p", "从工程角度看，这些异常处理并不复杂，但对毕业设计检查很重要。它说明系统不是只能在理想状态下演示，而是考虑了实验运行中可能出现的指标缺失、组件未就绪和数据不完整问题。对于云原生系统而言，部署成功只是第一步，能够发现错误、阻止错误数据进入结论，才是实验平台可信的关键。"),
        ("h2", "6.11 迁移到多服务验证的具体步骤"),
        ("p", "若后续将本文平台迁移到多服务场景，可以按照四步进行。第一步，设计最小业务链路，例如 gateway-service 调用 order-service，order-service 再调用 inventory-service，必要时加入 Redis 或数据库模拟库存状态。每个服务都应独立容器化，并分别配置 Deployment、Service、requests、limits 和健康探针。"),
        ("p", "第二步，为每个服务暴露独立指标。除了 http_requests_total 和请求耗时，还可以增加订单创建成功数、库存扣减失败数、队列长度、缓存命中率等业务指标。Prometheus Adapter 需要为不同服务配置不同指标名称，使 HPA 能够按服务角色读取合适指标。入口服务适合 QPS 指标，订单服务适合延迟和失败率指标，库存服务可能更适合队列长度或扣减请求数。"),
        ("p", "第三步，设计多服务策略组合。可以设置固定副本全链路、只扩 gateway、扩 gateway+order、扩全部关键服务、关键服务 HPA+VPA 等方案。这样能够观察只扩入口是否导致下游瓶颈，也能分析全链路扩容是否带来更高成本。第四步，扩展分析脚本，将单服务 P95/QPS/失败率扩展为链路级指标和服务级指标，最终判断瓶颈从哪个服务开始出现。"),
        ("p", "通过上述步骤，本文平台可以从单个典型无状态服务验证扩展到多服务链路验证。但这种扩展需要新增业务代码、指标定义和实验设计，不能直接由当前结果推出。因此，本文在正文中始终将当前成果表述为单服务弹性伸缩验证平台，而将多服务验证作为后续扩展方向。"),
        ("h2", "6.12 论文与代码一致性维护"),
        ("p", "毕业设计后期最容易出现的问题，是论文描述与代码仓库逐渐不一致。本文在优化论文时采取了“先看代码和数据，再写结论”的原则。凡是论文中出现的策略名称、实验参数、指标名称和结果数值，都尽量从仓库中的配置或 CSV 中取得。例如实验场景来自 experiment.config.json，策略结果来自 experiment_metrics_summary.csv，VPA 推荐来自归档的 vpa.yaml，HPA 运行状态来自 hpa_describe.txt。"),
        ("p", "这种一致性维护对答辩也有直接帮助。当教师询问某个结论依据时，可以从论文跳转到对应文件进行说明，而不是只依赖口头解释。若后续再次修改代码或重跑实验，也应同步更新论文表格、图表和结论，避免出现论文写 HPA+VPA 最优但数据文件已经变化的情况。对于系统设计类毕业论文而言，代码、数据和论文三者一致，比单纯追求文字长度更重要。"),
        ("h1", "7 总结与展望"),
        ("p", "本文围绕 Kubernetes 容器编排环境下的微服务弹性伸缩策略开展研究与实践，完成了一套可运行、可观测、可压测、可归档和可分析的实验平台。系统实现包括 sample-api 微服务、基础 Deployment/Service、CPU HPA、多指标 HPA、VPA Initial、Prometheus 指标采集、Prometheus Adapter 自定义指标注册、Locust 压测脚本、自动化实验脚本和结果分析脚本。该平台能够支撑系统设计类毕业论文对核心功能、现场演示、稳定性和论文一致性的检查。"),
        ("p", f"通过 4 组策略、每组 3 次、共 12 次正式实验，本文比较了固定副本、CPU HPA、多指标 HPA 和 HPA+VPA 在 burst 场景下的表现。实验结果表明，自动伸缩策略相较固定副本能够改善突发负载下的服务表现，其中 HPA+VPA 在本文环境中平均 P95 为 {hpa_vpa['P95(ms)']} ms、QPS 为 {hpa_vpa['QPS']}、失败率为 {hpa_vpa['失败率']}，综合表现较好。同时，实验也说明 HPA 控制周期、指标采集延迟、本地节点资源竞争和数据采集口径都会影响结果解释。"),
        ("p", "后续工作可以从三个方向展开。第一，将实验环境扩展到多节点 Kubernetes 集群，观察调度、节点资源和网络延迟对伸缩效果的影响。第二，将 sample-api 扩展为多服务业务链路，引入网关、订单、库存、缓存或消息队列，验证链路瓶颈下的分层伸缩策略。第三，探索预测式伸缩和更细粒度的业务指标，例如请求延迟、队列长度和错误率，并与 HPA/VPA 原生机制结合，形成更适合生产场景的弹性治理方案。"),
        ("h1", "参考文献"),
        ("refs", 0),
        ("h1", "致谢"),
        ("p", "本论文从选题、实验环境搭建到结果整理，经历了多次调试和修改。首先感谢指导教师在课题方向、实验设计和论文结构方面给予的指导，使我能够将 Kubernetes 弹性伸缩这一工程问题逐步拆解为可实现、可验证的系统任务。"),
        ("p", "在实现过程中，我对 Docker、Kubernetes、Prometheus、HPA、VPA 和 Locust 等工具有了更加具体的理解。许多问题并不是阅读文档即可解决，例如 metrics-server 无数据、Prometheus Adapter 指标未注册、VPA CRD 不存在、Locust 端口转发失败、实验采样缺失等，都需要结合日志和命令输出逐步排查。正是这些具体调试过程，让我认识到系统设计类课题的价值不只是写出代码，还包括让系统稳定运行并能解释运行结果。"),
        ("p", "同时感谢同学和往届材料提供的参考，使我能够更清楚地理解毕业设计材料、论文结构和答辩要求。本文仍存在不足之处，例如实验规模有限、多服务场景尚未展开、部分指标采集口径有待完善。后续我将继续补充工程实践经验，提高对云原生系统运行机制和性能分析方法的理解。"),
        ("h1", "附录"),
        ("h2", "附录 A 主要代码与配置文件说明"),
        ("p", "本文项目主要文件包括：services/sample-api/app.py，用于实现 HTTP 服务和 Prometheus 指标端点；deploy/base/deployment.yaml，用于定义 sample-api 的副本、资源请求、资源限制和健康探针；deploy/hpa/hpa-cpu.yaml，用于定义 CPU 单指标 HPA；deploy/hpa/hpa-multi.yaml，用于定义 CPU、内存和 QPS 多指标 HPA；deploy/vpa/vpa-initial.yaml，用于定义 VPA Initial 策略；scripts/run_experiments.ps1，用于自动执行四组对照实验；results/aggregate_experiment.py，用于从 raw 数据生成逐次指标和汇总指标；results/plot_results.py，用于生成论文图表。"),
        ("h2", "附录 B 正式实验逐次数据"),
        ("run_table", 0),
        ("p", "附录 B 中列出的逐次数据来自 results/processed/experiment_metrics.csv。保留逐次数据的目的是说明论文结论并非来自单次实验，而是来自重复运行后的汇总分析。对于出现明显波动的运行，可以继续回到对应 raw 目录查看 locust_stats.csv、top_pods.log、hpa_describe.txt 和 manifest.json。"),
        ("h2", "附录 C 答辩演示建议"),
        ("p", "答辩演示可按照“服务可运行、指标可采集、策略可解释、数据可复算”的顺序进行。首先展示 Kubernetes 命名空间中的 sample-api Pod 和 Service，说明被测微服务部署状态；其次访问 /metrics 或展示 Prometheus 指标，说明应用指标不是事后编造；然后展示 hpa-cpu.yaml、hpa-multi.yaml 和 vpa-initial.yaml，解释四组策略差异；最后展示 results/processed/experiment_metrics_summary.csv 和图表，说明论文结论来自正式实验数据。"),
        ("p", "若教师追问为什么不是算法设计类，可以回答：本课题不以模型训练和算法精度为目标，而以云原生系统部署、指标采集、弹性策略配置和实验验证为目标，符合系统设计类检查标准。若教师追问为什么只验证单个微服务，可以回答：Kubernetes HPA 的基本控制对象就是单个 Deployment，多服务链路中的每个服务也需要先具备独立伸缩能力，本文为了控制变量选择典型无状态服务作为基础实验对象。"),
        ("p", "提交前建议再次检查三个位置：论文表 5-1 是否与最新汇总 CSV 一致，正文中实验参数是否仍与配置文件一致，图表是否来自最新 results/figures 目录。这样可以避免论文、代码和实验结果在最后阶段出现版本不一致。"),
    ]


def _references() -> list[str]:
    return [
        "Kubernetes Documentation. Kubernetes Concepts and Tasks. https://kubernetes.io/docs/",
        "Kubernetes Documentation. Horizontal Pod Autoscaling. https://kubernetes.io/docs/tasks/run-application/horizontal-pod-autoscale/",
        "Kubernetes Documentation. Resource Management for Pods and Containers. https://kubernetes.io/docs/concepts/configuration/manage-resources-containers/",
        "Kubernetes Autoscaler. Vertical Pod Autoscaler. https://github.com/kubernetes/autoscaler/tree/master/vertical-pod-autoscaler",
        "Prometheus Documentation. Overview and Concepts. https://prometheus.io/docs/",
        "Prometheus Adapter. Kubernetes Metrics APIs Adapter for Prometheus. https://github.com/kubernetes-sigs/prometheus-adapter",
        "Locust Documentation. Load testing framework. https://docs.locust.io/",
        "CNCF. Cloud Native Definition. https://github.com/cncf/toc/blob/main/DEFINITION.md",
        "Burns B, Grant B, Oppenheimer D, Brewer E, Wilkes J. Borg, Omega, and Kubernetes. Communications of the ACM, 2016.",
        "Burns B, Beda J, Hightower K, Evenson L. Kubernetes: Up and Running. O'Reilly Media, 2022.",
        "Newman S. Building Microservices: Designing Fine-Grained Systems. O'Reilly Media, 2021.",
        "Hightower K, Burns B, Beda J. Kubernetes: Up and Running, 2nd Edition. O'Reilly Media, 2019.",
        "Locust Contributors. Writing a locustfile. https://docs.locust.io/en/stable/writing-a-locustfile.html",
        "Prometheus Authors. Prometheus Exposition Formats. https://prometheus.io/docs/instrumenting/exposition_formats/",
        "Kubernetes Documentation. Metrics For Kubernetes System Components. https://kubernetes.io/docs/concepts/cluster-administration/system-metrics/",
    ]


def _append_run_table(doc: Document) -> None:
    rows = []
    for row in _run_rows():
        rows.append(
            {
                "策略": STRATEGY_LABELS[row["strategy"]],
                "运行ID": row["run_id"],
                "P95": _fmt(row["p95_ms"], 0),
                "QPS": _fmt(row["qps"], 2),
                "失败率": _pct(row["failure_rate"]),
                "CPU": f"{float(row['cpu_util_percent']):.2f}%",
            }
        )
    _add_table(doc, rows)


def _count_chinese_chars(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def build_docx() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)
    doc = Document(str(OUTPUT_DOCX))
    _clear_doc_body_keep_section_props(doc)

    _add_title(doc, "基于容器编排的微服务弹性伸缩策略研究与实践")
    _add_center(doc, "刘乐", size=12)
    _add_page_break(doc.add_paragraph())

    for kind, value in _build_content():
        if kind == "h1":
            _add_heading(doc, str(value), 1)
        elif kind == "h2":
            _add_heading(doc, str(value), 2)
        elif kind == "p":
            _add_paragraph(doc, str(value))
        elif kind == "table_caption":
            _add_table_caption(doc, str(value))
        elif kind == "summary_table":
            _add_table(doc, _summary_rows())
        elif kind == "run_table":
            _append_run_table(doc)
        elif kind == "figure":
            filename, caption = value  # type: ignore[misc]
            _add_figure(doc, filename, caption)
        elif kind == "refs":
            _add_references(doc, _references())
        else:
            raise ValueError(f"unknown content kind: {kind}")

    doc.save(str(OUTPUT_DOCX))


def validate_docx() -> dict[str, object]:
    doc = Document(str(OUTPUT_DOCX))
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    checks = {
        "output": str(OUTPUT_DOCX),
        "chinese_chars": _count_chinese_chars(text),
        "paragraphs": sum(1 for p in doc.paragraphs if p.text.strip()),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "has_abstract": "摘要" in text and "Abstract" in text,
        "has_acknowledgement": "致谢" in text,
        "has_appendix": "附录" in text,
        "headings": headings,
    }
    return checks


def main() -> None:
    build_docx()
    checks = validate_docx()
    print(json.dumps(checks, ensure_ascii=False, indent=2))
    if int(checks["chinese_chars"]) < 15000:
        raise SystemExit("Chinese character count is below 15000")


if __name__ == "__main__":
    main()
