#!/usr/bin/env python3
"""Convert markdown thesis draft to DOCX using a sample DOCX style template."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


Token = Tuple[str, int, str]


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)\s*$")
ORDERED_LIST_RE = re.compile(r"^\s*\d+\.\s+(.+)$")
UNORDERED_LIST_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def _clean_inline_markdown(text: str) -> str:
    text = INLINE_CODE_RE.sub(r"\1", text)
    return text.strip()


def _parse_markdown(md_text: str) -> List[Token]:
    tokens: List[Token] = []
    paragraph_lines: List[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        paragraph_text = "\n".join(paragraph_lines).strip()
        if paragraph_text:
            tokens.append(("paragraph", 0, _clean_inline_markdown(paragraph_text)))
        paragraph_lines.clear()

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        if not line.strip():
            flush_paragraph()
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            text = _clean_inline_markdown(heading_match.group(2))
            tokens.append(("heading", level, text))
            continue

        ordered_list_match = ORDERED_LIST_RE.match(line)
        if ordered_list_match:
            flush_paragraph()
            tokens.append(("list", 1, _clean_inline_markdown(line.strip())))
            continue

        unordered_list_match = UNORDERED_LIST_RE.match(line)
        if unordered_list_match:
            flush_paragraph()
            tokens.append(("list", 0, _clean_inline_markdown(line.strip())))
            continue

        paragraph_lines.append(_clean_inline_markdown(line))

    flush_paragraph()
    return tokens


def _clear_doc_body_keep_section_props(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _resolve_style_name(doc: Document, preferred: str, fallback: str) -> str:
    style_names = {s.name for s in doc.styles}
    if preferred in style_names:
        return preferred
    return fallback


def _add_title(doc: Document, title: str) -> None:
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph("")


def _add_tokens(doc: Document, tokens: Sequence[Token]) -> None:
    h1_style = _resolve_style_name(doc, "Heading 1", "Normal")
    h2_style = _resolve_style_name(doc, "Heading 2", "Normal")
    h3_style = _resolve_style_name(doc, "Heading 3", "Normal")
    list_style = _resolve_style_name(doc, "List Paragraph", "Normal")
    normal_style = _resolve_style_name(doc, "Normal", "Normal")

    title_consumed = False
    for token_type, level, text in tokens:
        if not text:
            continue

        if token_type == "heading":
            if level == 1 and not title_consumed:
                _add_title(doc, text)
                title_consumed = True
                continue
            if level <= 2:
                style = h1_style
            elif level == 3:
                style = h2_style
            else:
                style = h3_style
            doc.add_paragraph(text, style=style)
            continue

        if token_type == "list":
            doc.add_paragraph(text, style=list_style)
            continue

        doc.add_paragraph(text, style=normal_style)


def convert_markdown_to_docx(md_path: Path, sample_docx_path: Path, output_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    tokens = _parse_markdown(md_text)

    doc = Document(str(sample_docx_path))
    _clear_doc_body_keep_section_props(doc)
    _add_tokens(doc, tokens)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--md", required=True, type=Path, help="Input markdown file")
    parser.add_argument(
        "--sample-docx", required=True, type=Path, help="Sample DOCX for style/template reference"
    )
    parser.add_argument("--out", required=True, type=Path, help="Output DOCX path")
    return parser


def main(argv: Iterable[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)

    convert_markdown_to_docx(
        md_path=args.md,
        sample_docx_path=args.sample_docx,
        output_path=args.out,
    )
    print(f"Generated: {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
